import os
import tempfile
import uuid
import time
from pathlib import Path

# Document processing
import PyPDF2
import docx
import pytesseract
from PIL import Image
import io

# Vector database and embeddings
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain_core.documents import Document

def extract_text_from_pdf(file):
    """Extract text from PDF using PyPDF2"""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page_num].extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    """Extract text from DOCX using python-docx"""
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_image(file):
    """Extract text from image using OCR"""
    image = Image.open(file)
    
    # Preprocess the image for better OCR results
    # Convert to grayscale
    image = image.convert('L')
    
    # Perform OCR
    text = pytesseract.image_to_string(image)
    
    return text

def extract_text_from_file(file, file_extension):
    """Extract text from file based on its extension"""
    if file_extension == ".pdf":
        return extract_text_from_pdf(file)
    elif file_extension in [".docx", ".doc"]:
        return extract_text_from_docx(file)
    elif file_extension in [".jpg", ".jpeg", ".png"]:
        return extract_text_from_image(file)
    elif file_extension == ".txt":
        return file.read().decode("utf-8")
    else:
        return None

def process_document(file, doc_type, case_id, doc_title, tags=[]):
    """Process a document and extract its text and metadata"""
    import streamlit as st
    import gc
    
    # Create a temporary file
    tmp_path = None
    
    try:
        # Create a named temporary file without auto-deletion
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.name).suffix) as tmp:
            tmp.write(file.getvalue())
            tmp_path = tmp.name
            # Make sure file is closed properly before proceeding
            tmp.flush()
            os.fsync(tmp.fileno())
        
        # Extract text based on file type
        file_extension = Path(file.name).suffix.lower()
        
        # Ensure the file exists before proceeding
        if not os.path.exists(tmp_path):
            st.error(f"Temporary file {tmp_path} does not exist.")
            return None
            
        # Force garbage collection to clean up any lingering file handles
        gc.collect()
        
        if file_extension == ".pdf":
            loader = PyPDFLoader(tmp_path)
            pages = loader.load_and_split()
            text = "\n".join([page.page_content for page in pages])
            
            # For OCR on PDF pages, use a safer approach
            try:
                # Try to use PyMuPDF for PDF page rendering and OCR
                try:
                    import fitz  # PyMuPDF
                    ocr_text = ""
                    pdf_document = fitz.open(tmp_path)
                    
                    for page_num in range(len(pdf_document)):
                        page = pdf_document[page_num]
                        # Convert page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2x zoom for better OCR
                        img_bytes = pix.pil_tobytes(format="PNG")
                        img = Image.open(io.BytesIO(img_bytes))
                        
                        # Perform OCR on the page image
                        page_text = pytesseract.image_to_string(img)
                        if page_text.strip():
                            ocr_text += f"\n\n[OCR from page {page_num+1}]:\n{page_text}"
                    
                    # Add OCR text to the document text if any was found
                    if ocr_text.strip():
                        text += "\n\n[OCR TEXT FROM PDF PAGES]\n" + ocr_text
                    
                    # Make sure to close the PDF document
                    pdf_document.close()
                        
                except ImportError:
                    # If PyMuPDF is not installed, fallback to simpler method
                    st.info("PyMuPDF (fitz) not installed. Using fallback OCR method for PDFs.")
                    # Use pdf2image as an alternative if available
                    try:
                        from pdf2image import convert_from_path
                        images = convert_from_path(tmp_path)
                        ocr_text = ""
                        
                        for i, image in enumerate(images):
                            page_text = pytesseract.image_to_string(image)
                            if page_text.strip():
                                ocr_text += f"\n\n[OCR from page {i+1}]:\n{page_text}"
                        
                        if ocr_text.strip():
                            text += "\n\n[OCR TEXT FROM PDF PAGES]\n" + ocr_text
                    except ImportError:
                        st.warning("Neither PyMuPDF nor pdf2image are installed. OCR for PDF pages skipped.")
                
            except Exception as e:
                st.warning(f"Error performing OCR on PDF: {str(e)}")
                    
        elif file_extension in [".docx", ".doc"]:
            loader = Docx2txtLoader(tmp_path)
            pages = loader.load_and_split()
            text = "\n".join([page.page_content for page in pages])
        elif file_extension == ".txt":
            loader = TextLoader(tmp_path)
            pages = loader.load_and_split()
            text = "\n".join([page.page_content for page in pages])
        elif file_extension in [".jpg", ".jpeg", ".png"]:
            text = extract_text_from_image(tmp_path)
            # Create a proper document object instead of just a string
            from langchain_core.documents import Document
            pages = [Document(page_content=text, metadata={"source": file.name, "page": 0})]
        else:
            st.error(f"Unsupported file format: {file_extension}")
            safely_delete_temp_file(tmp_path)
            return None
            
        # Clean up the temporary file
        safely_delete_temp_file(tmp_path)
        
        # Create document object
        doc_id = str(uuid.uuid4())
        document = {
            "id": doc_id,
            "title": doc_title or file.name,
            "type": doc_type,
            "case_id": case_id,
            "text": text,
            "filename": file.name,
            "uploaded_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "pages": pages,
            "tags": tags
        }
        
        return document
        
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        if tmp_path:
            safely_delete_temp_file(tmp_path)
        return None

def safely_delete_temp_file(tmp_path):
    """Safely delete a temporary file with error handling"""
    import streamlit as st
    
    if not tmp_path or not os.path.exists(tmp_path):
        return
        
    try:
        os.unlink(tmp_path)
    except PermissionError as e:
        st.warning(f"Could not delete temporary file {tmp_path}. It will be cleaned up later.")
        # Schedule file for deletion on Windows when process exits
        try:
            import atexit
            @atexit.register
            def delete_temp_file():
                try:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                except:
                    pass
        except:
            pass
    except Exception as e:
        st.warning(f"Error deleting temporary file: {str(e)}")

def parse_tags(tags_input):
    """Parse tags from comma-separated string"""
    if not tags_input:
        return []
    
    # Split by comma and clean whitespace
    tags = [tag.strip() for tag in tags_input.split(',')]
    
    # Remove empty tags
    tags = [tag for tag in tags if tag]
    
    return tags